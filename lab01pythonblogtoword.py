#!/usr/bin/env python
# coding: utf-8

# # 利用python爬取博客上单一作者的所有文章保存到本地
# 看上博客上一个作者的文章，想一次性下载到一个word文件中，并且可以设置好目录，通过word的“导航窗格”快速定位单篇文章。一劳永逸，从此再也不用去博客上一篇一篇地翻阅了。整理一下步骤：
# 1. 先获取到所有文章的标题、发表日期、链接http://blog.sina.com.cn/s/articlelist_5119330124_0_1.html
# 2. 通过链接获取文章的内容
# 3. 将文章标题作为“1级”，发表日期和内容作为正文写入word文件
# 4. 保存wrod文件
# 
# 下面就按照以上步骤进行操作。

# 先进入到目标博客的主页，点击“博文目录”，这样就在网址栏看到“blog.sina.com.cn/s/articlelist_5119330124_0_1.html” 。再点击下一页，可以看到网址末尾的“1”变成了“2”。这样我们就知道所有页对应的网址了（尾号从1到5）。

# 先挑第一页的网址，定位我们需要的信息，以便后续批量爬取。
# 
# 在博文的标题和发表日期上分别点右键，选择“检查”.
# 
# 可见博文标题和博文链接都位于`class="atc_title"`下面，发表时间位于`class="atc_tm SG_txtc"`下面。因此使用`soup.select('.atc_title')`就可以获取当前网页的所有博文的链接和标题；使用`soup.select('.atc_tm')`可获取所有博文的发表日期。

# In[2]:


import requests
from bs4 import BeautifulSoup
# 测试数据：http://blog.sina.com.cn/s/articlelist_5119330124_0_1.html
url=input("请输入博客作者的网页：")


# In[3]:


wb_data = requests.get(url)
soup = BeautifulSoup(wb_data.content)


# In[4]:


# 获取当页所有文章的标题和链接
soup.select('.atc_title')


# In[5]:


# 获取当页所有文章的发表时间
soup.select('.atc_tm')


# 如上获取的文章标题及链接信息是存在一个大列表中的。现在以第一个元素为例从中提取出链接和标题信息。观察发现链接位于`a`标签里的`href`里面，于是使用`select`方法选中`a`标签，可以看到结果是一个新的列表（如下）。

# In[6]:


soup.select('.atc_title')[0].select('a')


# 然后再从这个新列表中提取出链接和标题。使用`get("href")`方法获得链接；使用`text`方法获得标题。

# In[7]:


soup.select('.atc_title')[0].select('a')[0].get("href")


# In[8]:


soup.select('.atc_title')[0].select('a')[0].text


# 发表时间的获取就简单很多了，直接用`text`方法即可。

# In[9]:


soup.select('.atc_tm')[0].text


# 单页的信息搞定，然后就可以批量处理了。使用`for`循环遍历所有页，然后逐个提取。因为我们已知作者的文章共有5页，所以直接使用`range(1,6)`。将最终的信息存入字典`all_links`。其中，“标题”作为键，文章链接和发表时间作为值。通过`len(all_links)`查看获取的文章链接数，一共211篇文章。

# In[10]:


# 获取所有博客文章的链接
import requests
from bs4 import BeautifulSoup

all_links = {}
for i in range(1,6):
    wb_data = requests.get(url)
    soup = BeautifulSoup(wb_data.content)
    links = soup.select('.atc_title')
    times = soup.select('.atc_tm')
    for i in range(len(links)):
        http_link = links[i].select('a')[0].get('href')
        c="https:"
        link=c+http_link
        title = links[i].text.strip()
        time = times[i].text
        all_links[title] = [link, time] 


# In[11]:


len(all_links)


# In[12]:


all_links


# 拿到所有文章链接后，先取一个来测试一下如何获取页面的文字。在文字上点右键，选择“检查”，可见其内容位于`class=articalContent   newfont_family`里面，因此使用`soup.select(".articalContent.newfont_family")`就可以获取到（注意articalContent和newfont_family之间的空格要用"."代替）。将其存入`article`变量，显示一下，可以看到这是一个大列表，其中的文本就是我们需要的内容。下面就需要将文本单独提取出来。

# In[13]:


# 获取单篇文章中的文字
url = 'https://blog.sina.com.cn/s/blog_13122c74c0102zcw5.html'
wb_data = requests.get(url)
soup = BeautifulSoup(wb_data.content)
article = soup.select(".articalContent.newfont_family")
article


# 直接使用`text`方法就能提取出来。“\xa0”是个什么鬼？明显不是我们要的，百度了一下，说是什么不间断空格符。管他呢，直接使用`replace("\xa0","")`删掉，这下就美丽了。“\n”是换行，就不要删了，保持原格式比较好。

# In[14]:


article[0].text


# In[15]:


article[0].text.replace("\xa0","")


# In[16]:


# 获取单篇文章中的图片链接
wb_data = requests.get(url)
soup = BeautifulSoup(wb_data.content)
img_link = soup.select(".articalContent.newfont_family")[0].find_all("img")[0].get("real_src")


# In[21]:


soup.select(".articalContent.newfont_family")[0].find_all("img")[0].get("real_src")


# In[22]:


# 图片下载函数
def downloadImg(img_url, file_path):
    req = requests.get(url=img_url) 
    with open(file_path, 'wb') as f: 
        f.write(req.content)
downloadImg(url,'1.jpg')


# 以上理顺，就可以大刀阔斧地开干了。定义一个函数`to_word`，一个参数，就是上面获取到的数据字典`all_links`。设定好`header`，假装是浏览器在访问。然后新建一个word文档，设置全局字体为宋体。因为有些文章被加密，无法访问并获取内容，所以最终获取到的文章数不一定等于链接数。于是增加一个初始值为0的计数器，用于记录写入word文档中的文章数，以便心中有数。然后遍历所有文章的标题，将标题按照“1级”写入word文档，这样才能在“导航窗格”看到文章目录，方便后续选取阅读。日期和内容都作为段落写入。有些文章被加密，获取不到内容，此时`article`变量为空，所以加个if语句判断，以免程序崩溃。每写入一篇文章，计数器自动加1，然后通过`print`输出信息。最后保存文件，366页，35万字的博客就到手了，结果是美丽的！从此阅读博客文章轻松多了。

# In[23]:


# 写入标题，内容到word文件
import docx
from docx.oxml.ns import qn #用于应用中文字体

def to_word(all_links):
    header = {"User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/49.0.2623.221 Safari/537.36 SE 2.X MetaSr 1.0"}
    doc=docx.Document() #新建word文档
    doc.styles['Normal'].font.name=u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    counter = 0 #计数器，用于记录写入word的文章数
    for title in all_links.keys():
        doc.add_heading(title,1)
        date = all_links[title][1][:10]#只取日期，不要时间
        doc.add_paragraph(date)
        wb_data = requests.get(all_links[title][0],headers = header)
        soup = BeautifulSoup(wb_data.content)        
        article = soup.select(".articalContent.newfont_family")
        #有些文章被加密，获取不到内容，此时article为空，所以加个if语句判断
        if article:
            text = article[0].text.replace("\xa0","")
            doc.add_paragraph(text)
            print(f"写入文章 {title} 。")
            counter += 1
    print(f"共写入 {counter} 篇文章。")
    doc.save("文章.docx")  
    
to_word(all_links)
print('保存完成')

